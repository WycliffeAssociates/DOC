from typing import Mapping, Optional, Sequence

from document.config import settings
from document.domain.assembly_strategies.assembly_strategy_utils import (
    # commentary_book_intro,
    book_title,
    chapter_commentary,
    chapter_heading,
    chapter_intro,
    tn_chapter_verses,
    tq_chapter_verses,
)
from document.domain.assembly_strategies_docx.assembly_strategy_utils import (
    add_hr,
    create_docx_subdoc,
    add_one_column_section,
    add_two_column_section,
    add_page_break,
)

from document.domain.bible_books import BOOK_NAMES
from document.domain.model import (
    AssemblyLayoutEnum,
    BCBook,
    ChunkSizeEnum,
    LangDirEnum,
    TNBook,
    TQBook,
    TWBook,
    USFMBook,
)
from docx import Document  # type: ignore
from docxcompose.composer import Composer  # type: ignore

logger = settings.logger(__name__)


def assemble_content_by_lang_then_book(
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
    assembly_layout_kind: AssemblyLayoutEnum,
    chunk_size: ChunkSizeEnum,
    book_names: Mapping[str, str] = BOOK_NAMES,
) -> Composer:
    """
    Group content by language and then by book and then pass content
    and a couple other parameters, assembly_layout_kind and
    chunk_size, to interleaving strategy to do the actual
    interleaving.
    """
    book_id_map = dict((id, pos) for pos, id in enumerate(BOOK_NAMES.keys()))
    composers: list[Composer] = []
    most_lang_codes = max(
        [
            [usfm_book.lang_code for usfm_book in usfm_books],
            [tn_book.lang_code for tn_book in tn_books],
            [tq_book.lang_code for tq_book in tq_books],
            [tw_book.lang_code for tw_book in tw_books],
            [bc_book.lang_code for bc_book in bc_books],
        ],
        key=lambda x: len(x),
    )
    most_book_codes = max(
        [
            [usfm_book.book_code for usfm_book in usfm_books],
            [tn_book.book_code for tn_book in tn_books],
            [tq_book.book_code for tq_book in tq_books],
            [tw_book.book_code for tw_book in tw_books],
            [bc_book.book_code for bc_book in bc_books],
        ],
        key=lambda x: len(x),
    )
    for lang_code in most_lang_codes:
        for book_code in sorted(
            most_book_codes,
            key=lambda book_code: book_id_map[book_code],
        ):
            selected_usfm_books = [
                usfm_book
                for usfm_book in usfm_books
                if usfm_book.lang_code == lang_code and usfm_book.book_code == book_code
            ]
            usfm_book = selected_usfm_books[0] if selected_usfm_books else None
            usfm_book2 = (
                selected_usfm_books[1]
                if selected_usfm_books and len(selected_usfm_books) > 1
                else None
            )
            selected_tn_books = [
                tn_book
                for tn_book in tn_books
                if tn_book.lang_code == lang_code and tn_book.book_code == book_code
            ]
            tn_book = selected_tn_books[0] if selected_tn_books else None
            selected_tq_books = [
                tq_book
                for tq_book in tq_books
                if tq_book.lang_code == lang_code and tq_book.book_code == book_code
            ]
            tq_book = selected_tq_books[0] if selected_tq_books else None
            # TWBook doesn't really need to have a book_code attribute
            # because TW resources are language centric not book centric.
            # We could do something about that later if desired for
            # design cleanness sake.
            selected_tw_books = [
                tw_book
                for tw_book in tw_books
                if tw_book.lang_code == lang_code and tw_book.book_code == book_code
            ]
            tw_book = selected_tw_books[0] if selected_tw_books else None
            selected_bc_books = [
                bc_book
                for bc_book in bc_books
                if bc_book.lang_code == lang_code and bc_book.book_code == book_code
            ]
            bc_book = selected_bc_books[0] if selected_bc_books else None
        if usfm_book is not None:
            composers.append(
                assemble_usfm_by_book(
                    usfm_book,
                    tn_book,
                    tq_book,
                    tw_book,
                    usfm_book2,
                    bc_book,
                )
            )
        elif usfm_book is None and tn_book is not None:
            composers.append(
                assemble_tn_by_book(
                    usfm_book,
                    tn_book,
                    tq_book,
                    tw_book,
                    usfm_book2,
                    bc_book,
                )
            )
        elif usfm_book is None and tn_book is None and tq_book is not None:
            composers.append(
                assemble_tq_by_book(
                    usfm_book,
                    tn_book,
                    tq_book,
                    tw_book,
                    usfm_book2,
                    bc_book,
                )
            )
        elif (
            usfm_book is None
            and tn_book is None
            and tq_book is None
            and (tw_book is not None or bc_book is not None)
        ):
            composers.append(
                assemble_tw_by_book(
                    usfm_book,
                    tn_book,
                    tq_book,
                    tw_book,
                    usfm_book2,
                    bc_book,
                )
            )
    first_composer = composers[0]
    for composer in composers[1:]:
        first_composer.append(composer.doc)
    return first_composer


def assemble_usfm_by_book(
    usfm_book: Optional[USFMBook],
    tn_book: Optional[TNBook],
    tq_book: Optional[TQBook],
    tw_book: Optional[TWBook],
    usfm_book2: Optional[USFMBook],
    bc_book: Optional[BCBook],
) -> Composer:
    """
    Construct the HTML for a 'by book' strategy wherein at least
    usfm_book_content_unit exists.
    """
    doc = Document()
    composer = Composer(doc)
    if tn_book:
        if tn_book.book_intro:
            subdoc = create_docx_subdoc(
                tn_book.book_intro,
                tn_book.lang_code,
                tn_book and tn_book.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)
    if bc_book:
        if bc_book.book_intro:
            subdoc = create_docx_subdoc(
                bc_book.book_intro,
                bc_book.lang_code,
            )
            composer.append(subdoc)
    if usfm_book:
        # fmt: off
        is_rtl = usfm_book and usfm_book.lang_direction == LangDirEnum.RTL
        # fmt: on
        subdoc = create_docx_subdoc(
            book_title(usfm_book.book_code),
            usfm_book.lang_code,
            is_rtl,
        )
        composer.append(subdoc)
        for (
            chapter_num,
            chapter,
        ) in usfm_book.chapters.items():
            add_one_column_section(doc)
            tn_verses: str = ""
            tq_verses: str = ""
            chapter_intro_ = ""
            chapter_commentary_ = ""
            if tn_book:
                chapter_intro_ = chapter_intro(tn_book, chapter_num)
                tn_verses = tn_chapter_verses(tn_book, chapter_num)
            if bc_book:
                chapter_commentary_ = chapter_commentary(bc_book, chapter_num)
            if tq_book:
                tq_verses = tq_chapter_verses(tq_book, chapter_num)
            subdoc = create_docx_subdoc(
                chapter.content,
                usfm_book.lang_code,
                is_rtl,
            )
            composer.append(subdoc)
            if chapter_intro_:
                subdoc = create_docx_subdoc(chapter_intro_, usfm_book.lang_code, is_rtl)
                composer.append(subdoc)
            if chapter_commentary_:
                subdoc = create_docx_subdoc(
                    chapter_commentary_, usfm_book.lang_code, is_rtl
                )
                composer.append(subdoc)
            if tn_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tn_verses,
                    usfm_book.lang_code,
                    is_rtl,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tq_verses,
                    usfm_book.lang_code,
                    is_rtl,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            # TODO Get feedback on whether we should allow a user to select a primary _and_
            # a secondary USFM resource. If we want to limit the user to only one USFM per
            # document then we would want to control that in the UI and maybe also at the API
            # level. The API level control would be implemented in the DocumentRequest
            # validation.
            if usfm_book2:
                add_one_column_section(doc)
                # Here we add the whole chapter's worth of verses for the secondary usfm
                subdoc = create_docx_subdoc(
                    usfm_book2.chapters[chapter_num].content,
                    usfm_book.lang_code,
                    usfm_book2 and usfm_book2.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
            add_page_break(doc)
    return composer


def assemble_tn_by_book(
    usfm_book: Optional[USFMBook],
    tn_book: Optional[TNBook],
    tq_book: Optional[TQBook],
    tw_book: Optional[TWBook],
    usfm_book2: Optional[USFMBook],
    bc_book: Optional[BCBook],
) -> Composer:
    """
    Construct the HTML for a 'by book' strategy wherein at least
    tn_book exists.
    """
    doc = Document()
    composer = Composer(doc)
    if tn_book:
        if tn_book.book_intro:
            subdoc = create_docx_subdoc(
                tn_book.book_intro,
                tn_book.lang_code,
                tn_book and tn_book.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)
        if bc_book:
            if bc_book.book_intro:
                subdoc = create_docx_subdoc(
                    bc_book.book_intro,
                    tn_book.lang_code,
                )
                composer.append(subdoc)
        for chapter_num in tn_book.chapters:
            add_one_column_section(doc)
            one_column_html = []
            one_column_html.append(chapter_heading(chapter_num))
            one_column_html.append(chapter_intro(tn_book, chapter_num))
            one_column_html_ = "".join(one_column_html)
            if one_column_html_:
                subdoc = create_docx_subdoc(
                    one_column_html_,
                    tn_book.lang_code,
                    tn_book and tn_book.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
            if bc_book:
                subdoc = create_docx_subdoc(
                    chapter_commentary(bc_book, chapter_num),
                    bc_book.lang_code,
                )
                composer.append(subdoc)
            tn_verses = tn_chapter_verses(tn_book, chapter_num)
            if tn_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tn_verses,
                    tn_book.lang_code,
                    tn_book and tn_book.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            tq_verses = tq_chapter_verses(tq_book, chapter_num)
            if tq_book and tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tq_verses,
                    tq_book.lang_code,
                    tq_book and tq_book.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            add_page_break(doc)
    return composer


def assemble_tq_by_book(
    usfm_book: Optional[USFMBook],
    tn_book: Optional[TNBook],
    tq_book: Optional[TQBook],
    tw_book: Optional[TWBook],
    usfm_book2: Optional[USFMBook],
    bc_book: Optional[BCBook],
) -> Composer:
    """
    Construct the HTML for a 'by book' strategy wherein at least
    tq_book exists.
    """
    doc = Document()
    composer = Composer(doc)
    if tq_book:
        for chapter_num in tq_book.chapters:
            add_one_column_section(doc)
            if bc_book:
                subdoc = create_docx_subdoc(
                    chapter_commentary(bc_book, chapter_num),
                    bc_book.lang_code,
                )
                composer.append(subdoc)
            subdoc = create_docx_subdoc(
                chapter_heading(chapter_num),
                tq_book.lang_code,
                tq_book and tq_book.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)
            tq_verses = tq_chapter_verses(tq_book, chapter_num)
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tq_verses,
                    tq_book.lang_code,
                    tq_book and tq_book.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
            add_page_break(doc)
    return composer


def assemble_tw_by_book(
    usfm_book: Optional[USFMBook],
    tn_book: Optional[TNBook],
    tq_book: Optional[TQBook],
    tw_book: Optional[TWBook],
    usfm_book2: Optional[USFMBook],
    bc_book: Optional[BCBook],
) -> Composer:

    """
    TW is handled outside this module, that is why no
    code for TW is explicitly included here.
    """
    doc = Document()
    composer = Composer(doc)
    if bc_book:
        subdoc = create_docx_subdoc(bc_book.book_intro, bc_book.lang_code)
        composer.append(subdoc)
        for chapter in bc_book.chapters.values():
            subdoc = create_docx_subdoc(chapter.commentary, bc_book.lang_code)
            composer.append(subdoc)
            add_page_break(doc)
    return composer
