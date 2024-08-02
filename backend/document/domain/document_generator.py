"""
Entrypoint for backend. Here incoming document requests are processed
and eventually a final document produced.
"""

import re
import shutil
import smtplib
import subprocess
import time
from collections import defaultdict
from datetime import datetime
from email.encoders import encode_base64
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from os.path import basename, exists, join
from pathlib import Path
from typing import Any, Mapping, Optional, Sequence, cast

import jinja2
from celery import current_task
from document.config import settings
from document.domain import exceptions, parsing, resource_lookup, worker
from document.domain.bible_books import BOOK_NAMES

from document.domain.assembly_strategies.assembly_strategies_book_then_lang_by_chapter import (
    assemble_content_by_book_then_lang,
)
from document.domain.assembly_strategies.assembly_strategies_lang_then_book_by_chapter import (
    assemble_content_by_lang_then_book,
)
from document.domain.assembly_strategies_docx import (
    assembly_strategies_book_then_lang_by_chapter as asd_book_then_lang,
)
from document.domain.assembly_strategies_docx import (
    assembly_strategies_lang_then_book_by_chapter as asd_lang_then_book,
)
from document.domain.assembly_strategies_docx.assembly_strategy_utils import add_hr
from document.domain.model import (
    AssemblyLayoutEnum,
    AssemblyStrategyEnum,
    Attachment,
    BCBook,
    ChunkSizeEnum,
    DocumentRequest,
    DocumentRequestSourceEnum,
    ResourceLookupDto,
    ResourceRequest,
    TNBook,
    TQBook,
    TWBook,
    TWNameContentPair,
    TWUse,
    USFMBook,
)
from document.utils.file_utils import file_needs_update, write_file
from docx import Document  # type: ignore
from docx.enum.section import WD_SECTION  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docxcompose.composer import Composer  # type: ignore
from docxtpl import DocxTemplate  # type: ignore
from htmldocx import HtmlToDocx  # type: ignore
from pydantic import Json

logger = settings.logger(__name__)


TEMPLATE_PATHS_MAP: Mapping[str, str] = {
    "book_intro": "backend/templates/tn/book_intro_template.md",
    "header_enclosing": "backend/templates/html/header_enclosing.html",
    "header_enclosing_landscape": "backend/templates/html/header_enclosing_landscape.html",  # used by dft project
    "header_no_css_enclosing": "backend/templates/html/header_no_css_enclosing.html",
    "header_compact_enclosing": "backend/templates/html/header_compact_enclosing.html",
    "footer_enclosing": "backend/templates/html/footer_enclosing.html",
    "cover": "backend/templates/html/cover.html",
    "email-html": "backend/templates/html/email.html",
    "email": "backend/templates/text/email.txt",
}


def contains_tw(resource_request: ResourceRequest, tw_regex: str = "tw.*") -> bool:
    """Return True if the resource_request describes a TW resource."""
    value = bool(re.compile(tw_regex).match(resource_request.resource_type))
    logger.debug(
        "resource_request: %s tests %s for TW resource type",
        resource_request,
        value,
    )
    return value


def document_request_key(
    resource_requests: Sequence[ResourceRequest],
    assembly_strategy_kind: AssemblyStrategyEnum,
    assembly_layout_kind: AssemblyLayoutEnum,
    chunk_size: ChunkSizeEnum,
    limit_words: bool,
    max_filename_len: int = 240,
    underscore: str = "_",
    hyphen: str = "-",
) -> str:
    """
    Create and return the document_request_key. The
    document_request_key uniquely identifies a document request.

    If the document request key is max_filename_len or more characters
    in length, then switch to using a shorter string that is based on the
    current time. The reason for this is that the document request key is
    used as the file name (with suffix appended) and each OS has a limit
    to how long a file name may be. max_filename_len should make room for
    the file suffix, e.g., ".html", to be appended.

    It is really useful to have filenames with semantic meaning and so
    those are preferred when possible, i.e., when the file name is not
    too long.
    """
    resource_request_keys = underscore.join(
        [
            hyphen.join(
                [
                    resource_request.lang_code,
                    resource_request.resource_type,
                    resource_request.book_code,
                ]
            )
            for resource_request in resource_requests
        ]
    )
    if any(contains_tw(resource_request) for resource_request in resource_requests):
        document_request_key = "{}_{}_{}_{}_{}".format(
            resource_request_keys,
            assembly_strategy_kind.value,
            assembly_layout_kind.value,
            chunk_size.value,
            "ltwt" if limit_words else "ltwf",
        )
    else:
        document_request_key = "{}_{}_{}_{}".format(
            resource_request_keys,
            assembly_strategy_kind.value,
            assembly_layout_kind.value,
            chunk_size.value,
        )
    if len(document_request_key) >= max_filename_len:
        # Likely the generated filename was too long for the OS where this is
        # running. In that case, use the current time as a document_request_key
        # value as doing so results in an acceptably short length.
        timestamp_components = str(time.time()).split(".")
        return "{}_{}".format(timestamp_components[0], timestamp_components[1])
    else:
        # Use the semantic filename which declaratively describes the
        # document request components.
        return document_request_key


def template_path(
    key: str, template_paths_map: Mapping[str, str] = TEMPLATE_PATHS_MAP
) -> str:
    """
    Return the path to the requested template give a lookup key.
    Return a different path if the code is running inside the Docker
    container.
    """
    return template_paths_map[key]


def template(template_lookup_key: str) -> str:
    """Return template as string."""
    with open(template_path(template_lookup_key), "r") as filepath:
        template = filepath.read()
    return template


def instantiated_email_template(document_request_key: str) -> str:
    """
    Instantiate Jinja2 template. Return instantiated template as string.
    """
    with open(template_path("email"), "r") as filepath:
        template = filepath.read()
    env = jinja2.Environment(autoescape=True).from_string(template)
    return env.render(data=document_request_key)


def instantiated_html_header_template(
    template_lookup_key: str, title1: str, title2: str, title3: str
) -> str:
    """
    Instantiate Jinja2 template. Return instantiated template as string.
    """
    template = ""
    with open(template_path(template_lookup_key), "r") as filepath:
        template = filepath.read()
    env = jinja2.Environment(autoescape=True).from_string(template)
    timestring = datetime.now().ctime()
    return env.render(
        timestring=timestring, title1=title1, title2=title2, title3=title3
    )


def enclose_html_content(
    content: str,
    document_html_header: str,
    document_html_footer: str = template("footer_enclosing"),
) -> str:
    """
    Write the enclosing HTML header and footer elements around the
    HTML body content for the document.
    """
    return "{}{}{}".format(document_html_header, content, document_html_footer)


def document_html_header(
    assembly_layout_kind: Optional[AssemblyLayoutEnum],
    generate_docx: bool,
    title1: str,
    title2: str,
    title3: str,
) -> str:
    """
    Choose the appropriate HTML header given the
    assembly_layout_kind. The HTML header, naturally, contains the CSS
    definitions and they in turn can be used to affect visual
    compactness.
    """
    if generate_docx:
        return template("header_no_css_enclosing")
    if assembly_layout_kind and assembly_layout_kind in [
        AssemblyLayoutEnum.ONE_COLUMN_COMPACT,
        AssemblyLayoutEnum.TWO_COLUMN_SCRIPTURE_LEFT_SCRIPTURE_RIGHT_COMPACT,
    ]:
        return instantiated_html_header_template(
            "header_compact_enclosing", title1, title2, title3
        )
    return instantiated_html_header_template("header_enclosing", title1, title2, title3)


# def uses_section(
#     uses: Sequence[TWUse],
#     translation_word_verse_section_header_str: str = settings.TRANSLATION_WORD_VERSE_SECTION_HEADER_STR,
#     unordered_list_begin_str: str = settings.UNORDERED_LIST_BEGIN_STR,
#     translation_word_verse_ref_item_fmt_str: str = settings.TRANSLATION_WORD_VERSE_REF_ITEM_FMT_STR,
#     unordered_list_end_str: str = settings.UNORDERED_LIST_END_STR,
#     book_numbers: Mapping[str, str] = bible_books.BOOK_NUMBERS,
#     book_names: Mapping[str, str] = bible_books.BOOK_NAMES,
#     num_zeros: int = 3,
# ) -> str:
#     """
#     Construct and return the 'Uses:' section which comes at the end of
#     a translation word definition and wherein each item points to
#     verses (as targeted by lang_code, book_id, chapter_num, and
#     verse_num) wherein the word occurs.
#     """
#     html: list[str] = []
#     html.append(translation_word_verse_section_header_str)
#     html.append(unordered_list_begin_str)
#     for use in uses:
#         html_content_str = translation_word_verse_ref_item_fmt_str.format(
#             use.lang_code,
#             book_numbers[use.book_id].zfill(num_zeros),
#             str(use.chapter_num).zfill(num_zeros),
#             str(use.verse_num).zfill(num_zeros),
#             book_names[use.book_id],
#             use.chapter_num,
#             use.verse_num,
#         )
#         html.append(html_content_str)
#     html.append(unordered_list_end_str)
#     return "\n".join(html)


def translation_words_section(
    tw_book: TWBook,
    usfm_books: Optional[Sequence[USFMBook]],
    limit_words: bool,
    resource_requests: Sequence[ResourceRequest],
    resource_type_name_fmt_str: str = settings.RESOURCE_TYPE_NAME_FMT_STR,
) -> str:
    """
    Build and return the translation words definition section, i.e.,
    the list of all translation words for this language, book combination.
    Limit the translation words to only those that appear in the USFM
    resouce chosen if limit_words is True and a USFM resource was also
    chosen.
    """

    content = []
    if tw_book.name_content_pairs:
        content.append(resource_type_name_fmt_str.format(tw_book.resource_type_name))
    selected_name_content_pairs = get_selected_name_content_pairs(
        tw_book, usfm_books, limit_words, resource_requests
    )
    for name_content_pair in selected_name_content_pairs:
        content.append(name_content_pair_content(name_content_pair, tw_book))
    return "".join(content)


def get_selected_name_content_pairs(
    tw_book: TWBook,
    usfm_books: Optional[Sequence[USFMBook]],
    limit_words: bool,
    resource_requests: Sequence[ResourceRequest],
    usfm_resource_types: Sequence[str] = settings.USFM_RESOURCE_TYPES,
) -> list[TWNameContentPair]:
    selected_name_content_pairs = []
    if usfm_books and limit_words:
        selected_name_content_pairs = filter_name_content_pairs(tw_book, usfm_books)
    elif not usfm_books and limit_words:
        usfm_books = fetch_usfm_book_content_units(resource_requests)
        selected_name_content_pairs = filter_name_content_pairs(tw_book, usfm_books)
    else:
        selected_name_content_pairs = tw_book.name_content_pairs
    return selected_name_content_pairs


def filter_name_content_pairs(
    tw_book: TWBook, usfm_books: Optional[Sequence[USFMBook]]
) -> list[TWNameContentPair]:
    selected_name_content_pairs = []
    if usfm_books:
        for name_content_pair in tw_book.name_content_pairs:
            for usfm_book in usfm_books:
                for chapter in usfm_book.chapters.values():
                    if re.search(
                        re.escape(name_content_pair.localized_word),
                        chapter.content,
                    ):
                        selected_name_content_pairs.append(name_content_pair)
                        break
    return selected_name_content_pairs


def fetch_usfm_book_content_units(
    resource_requests: Sequence[ResourceRequest],
    usfm_resource_types: Sequence[str] = settings.USFM_RESOURCE_TYPES,
) -> list[USFMBook]:
    usfm_resource_lookup_dtos = []
    for resource_request in resource_requests:
        for usfm_type in usfm_resource_types:
            resource_lookup_dto = resource_lookup.resource_lookup_dto(
                resource_request.lang_code,
                usfm_type,
                resource_request.book_code,
            )
            if resource_lookup_dto:
                usfm_resource_lookup_dtos.append(resource_lookup_dto)
    # Determine which resource URLs were actually found.
    found_usfm_resource_lookup_dtos = [
        resource_lookup_dto
        for resource_lookup_dto in usfm_resource_lookup_dtos
        if resource_lookup_dto.url is not None
    ]
    current_task.update_state(state="Provisioning USFM asset files for TW resource")
    t0 = time.time()
    resource_dirs = [
        resource_lookup.provision_asset_files(dto)
        for dto in found_usfm_resource_lookup_dtos
    ]
    t1 = time.time()
    logger.debug(
        "Time to provision USFM asset files (acquire and write to disk) for TW resource: %s",
        t1 - t0,
    )
    current_task.update_state(state="Parsing USFM asset files for TW resource")
    # Initialize found resources from their provisioned assets.
    usfm_book_content_units = [
        parsing.usfm_book_content(
            resource_lookup_dto,
            resource_dir,
            resource_requests,
            False,
        )
        for resource_lookup_dto, resource_dir in zip(
            found_usfm_resource_lookup_dtos, resource_dirs
        )
    ]
    return usfm_book_content_units


def name_content_pair_content(
    name_content_pair: TWNameContentPair,
    tw_book: TWBook,
    # include_uses_section: bool,
) -> str:
    name_content_pair.content = modify_content_for_anchors(name_content_pair, tw_book)
    # uses_section_ = ""
    # if (
    #     include_uses_section
    #     and name_content_pair.localized_word in book_content_unit.uses
    # ):
    #     uses_section_ = uses_section(
    #         book_content_unit.uses[name_content_pair.localized_word]
    #     )
    #     name_content_pair.content = f"{name_content_pair.content}{uses_section_}"
    return name_content_pair.content


def modify_content_for_anchors(
    name_content_pair: TWNameContentPair,
    book_content_unit: TWBook,
    opening_h3_fmt_str: str = settings.OPENING_H3_FMT_STR,
    opening_h3_with_id_fmt_str: str = settings.OPENING_H3_WITH_ID_FMT_STR,
) -> str:
    return name_content_pair.content.replace(
        opening_h3_fmt_str.format(name_content_pair.localized_word),
        opening_h3_with_id_fmt_str.format(
            book_content_unit.lang_code,
            name_content_pair.localized_word,
            name_content_pair.localized_word,
        ),
    )


def assemble_content(
    document_request_key: str,
    document_request: DocumentRequest,
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
    found_resource_lookup_dtos: Sequence[ResourceLookupDto],
) -> str:
    """
    Assemble and return the content from all requested resources according to the
    assembly_strategy requested.
    """
    t0 = time.time()
    content = ""
    if (
        document_request.assembly_strategy_kind
        == AssemblyStrategyEnum.LANGUAGE_BOOK_ORDER
    ):
        content = "".join(
            assemble_content_by_lang_then_book(
                usfm_books,
                tn_books,
                tq_books,
                tw_books,
                bc_books,
                cast(AssemblyLayoutEnum, document_request.assembly_layout_kind),
            )
        )
    elif (
        document_request.assembly_strategy_kind
        == AssemblyStrategyEnum.BOOK_LANGUAGE_ORDER
    ):
        content = "".join(
            assemble_content_by_book_then_lang(
                usfm_books,
                tn_books,
                tq_books,
                tw_books,
                bc_books,
                cast(AssemblyLayoutEnum, document_request.assembly_layout_kind),
            )
        )
    t1 = time.time()
    logger.debug("Time for interleaving document: %s", t1 - t0)
    t0 = time.time()
    # Add the translation words definition section for each language requested.
    unique_lang_codes = set()
    for tw_book in tw_books:
        if tw_book.lang_code not in unique_lang_codes:
            unique_lang_codes.add(tw_book.lang_code)
            content = "{}{}<hr/>".format(
                content,
                translation_words_section(
                    tw_book,
                    usfm_books,
                    document_request.limit_words,
                    document_request.resource_requests,
                ),
            )
    t1 = time.time()
    logger.debug("Time for add TW content to document: %s", t1 - t0)
    return content


def create_title_page_and_wrap_in_template(
    content: str,
    document_request: DocumentRequest,
    found_resource_lookup_dtos: Sequence[ResourceLookupDto],
) -> str:

    title1, title2 = get_languages_title_page_strings(found_resource_lookup_dtos)
    title3 = "Formatted for Translators"
    header = document_html_header(
        document_request.assembly_layout_kind,
        document_request.generate_docx,
        title1,
        title2,
        title3,
    )
    content = enclose_html_content(content, document_html_header=header)
    return content


def filter_unique_by_lang_code(tw_books: Sequence[TWBook]) -> list[TWBook]:
    unique_tw_books = []
    seen_lang_codes = set()
    for tw_book in tw_books:
        lang_code = tw_book.lang_code
        if lang_code not in seen_lang_codes:
            seen_lang_codes.add(lang_code)
            unique_tw_books.append(tw_book)
    return unique_tw_books


def assemble_docx_content(
    document_request_key: str,
    document_request: DocumentRequest,
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
) -> Composer:
    """
    Assemble and return the content from all requested resources according to the
    assembly_strategy requested.
    """
    t0 = time.time()
    composer = None
    if (
        document_request.assembly_strategy_kind
        == AssemblyStrategyEnum.LANGUAGE_BOOK_ORDER
    ):
        composer = asd_lang_then_book.assemble_content_by_lang_then_book(
            usfm_books,
            tn_books,
            tq_books,
            tw_books,
            bc_books,
            cast(AssemblyLayoutEnum, document_request.assembly_layout_kind),
            document_request.chunk_size,
        )
    elif (
        document_request.assembly_strategy_kind
        == AssemblyStrategyEnum.BOOK_LANGUAGE_ORDER
    ):
        composer = asd_book_then_lang.assemble_content_by_book_then_lang(
            usfm_books,
            tn_books,
            tq_books,
            tw_books,
            bc_books,
            cast(AssemblyLayoutEnum, document_request.assembly_layout_kind),
            document_request.chunk_size,
        )
    t1 = time.time()
    logger.debug("Time for interleaving document: %s", t1 - t0)
    tw_subdocs = []
    if tw_books:
        html_to_docx = HtmlToDocx()
        t0 = time.time()
        # Add the translation words definition section for each language requested.
        unique_tw_books = filter_unique_by_lang_code(tw_books)
        for tw_book in unique_tw_books:
            tw_subdoc = html_to_docx.parse_html_string(
                translation_words_section(
                    tw_book,
                    usfm_books,
                    document_request.limit_words,
                    document_request.resource_requests,
                )
            )
            if tw_subdoc.paragraphs:
                p = tw_subdoc.paragraphs[-1]
                add_hr(p)
                tw_subdocs.append(tw_subdoc)
        t1 = time.time()
        logger.debug("Time for adding TW content to document: %s", t1 - t0)
    # Now add any TW subdocs to the composer
    for tw_subdoc_ in tw_subdocs:
        composer.append(tw_subdoc_)
    return composer


def should_send_email(
    # email_address comes in as pydantic.EmailStr and leaves
    # the pydantic class validator as a str.
    email_address: Optional[str],
    send_email: bool = settings.SEND_EMAIL,
) -> bool:
    """
    Return True if configuration is set to send email and the user
    has supplied an email address.
    """
    return send_email and email_address is not None


def send_email_with_attachment(
    # email_address comes in as pydantic.EmailStr and leaves
    # the pydantic class validator as a str.
    email_address: Optional[str],
    attachments: list[Attachment],
    document_request_key: str,
    content_disposition: str = "attachment",
    from_email_address: str = settings.FROM_EMAIL_ADDRESS,
    smtp_password: str = settings.SMTP_PASSWORD,
    email_send_subject: str = settings.EMAIL_SEND_SUBJECT,
    smtp_host: str = settings.SMTP_HOST,
    smtp_port: int = settings.SMTP_PORT,
    comma_space: str = ", ",
) -> None:
    """
    If environment configuration allows sending of
    email, then send an email to the document request
    recipient's email with the document attached.
    """
    if email_address:
        sender = from_email_address
        email_password = smtp_password
        recipients = [email_address]
        logger.debug("Email sender %s, recipients: %s", sender, recipients)
        # Create the enclosing (outer) message
        outer = MIMEMultipart()
        outer["Subject"] = email_send_subject
        outer["To"] = comma_space.join(recipients)
        outer["From"] = sender
        # Add the attachments to the message
        for attachment in attachments:
            try:
                with open(attachment.filepath, "rb") as fp:
                    msg = MIMEBase(attachment.mime_type[0], attachment.mime_type[1])
                    msg.set_payload(fp.read())
                encode_base64(msg)
                msg.add_header(
                    "Content-Disposition",
                    content_disposition,
                    filename=basename(attachment.filepath),
                )
                outer.attach(msg)
            except Exception:
                logger.exception(
                    "Unable to open one of the attachments. Caught exception: "
                )
        # Get the email body
        message_body = instantiated_email_template(document_request_key)
        logger.debug("instantiated email template: %s", message_body)
        outer.attach(MIMEText(message_body, "plain"))
        composed = outer.as_string()
        # Send the email
        try:
            with smtplib.SMTP(smtp_host, smtp_port) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.ehlo()
                smtp.login(sender, email_password)
                smtp.sendmail(sender, recipients, composed)
                smtp.close()
            logger.info("Email sent!")
        except Exception:
            logger.exception("Unable to send the email. Caught exception: ")


# HTML to PDF converters:
# princexml ($$$$) (fastest) or same through docraptor ($$) but slow,
# wkhtmltopdf via pdfkit (can't handle column-count directive so can't use due to
# multi-column layouts we want; also no longer maintained),
# weasyprint (does a nice job, we use this),
# pagedjs-cli (does a really nice job, but is really slow - uses puppeteer underneath),
# electron-pdf (similar speed to wkhtmltopdf) which uses chrome underneath the hood,
# gotenburg which uses chrome under the hood and provides a nice api in Docker (untested),
# raw chrome headless (works well and is about the same speed as weasyprint),
# ebook-convert (currently blows up because docker runs chrome as root)
def convert_html_to_pdf(
    html_filepath: str,
    pdf_filepath: str,
    document_request_key: str,
) -> None:
    """
    Generate PDF from HTML and copy it to output directory.
    """
    assert exists(html_filepath)
    logger.info("Generating PDF %s...", pdf_filepath)
    t0 = time.time()
    weasyprint_command = "weasyprint {} {}".format(html_filepath, pdf_filepath)
    logger.debug("Generate PDF command: %s", weasyprint_command)
    subprocess.call(weasyprint_command, shell=True)
    t1 = time.time()
    logger.debug("Time for converting HTML to PDF: %s", t1 - t0)


# HTML to ePub converters:
# pandoc (this doesn't respect two column),
# html-to-epub which is written in go (this doesn't respect two column),
# ebook-convert (this respects two column).
def convert_html_to_epub(
    html_filepath: str,
    epub_filepath: str,
    document_request_key: str,
) -> None:
    """Generate ePub from HTML and copy it to output directory."""
    assert exists(html_filepath)
    ebook_convert_command = (
        "/calibre-bin/calibre/ebook-convert {} {} --no-default-epub-cover".format(
            html_filepath, epub_filepath
        )
    )
    logger.debug("Generate ePub command: %s", ebook_convert_command)
    t0 = time.time()
    subprocess.call(ebook_convert_command, shell=True)
    t1 = time.time()
    logger.debug("Time for converting HTML to ePub: %s", t1 - t0)


def convert_html_to_docx(
    html_filepath: str,
    docx_filepath: str,
    composer: Composer,
    layout_for_print: bool,
    title1: str = "title1",
    title2: str = "title2",
    title3: str = "Formatted for Translators",
    docx_template_path: str = settings.DOCX_TEMPLATE_PATH,
    docx_compact_template_path: str = settings.DOCX_COMPACT_TEMPLATE_PATH,
) -> None:
    """Generate Docx and copy it to output directory."""
    t0 = time.time()
    # Get data for front page of Docx template.
    title1 = title1
    title2 = title2
    title3 = title3
    # fmt: off
    template_path = docx_compact_template_path if layout_for_print else docx_template_path
    # fmt: on
    doc = DocxTemplate(template_path)
    toc_path = generate_docx_toc(docx_filepath)
    toc = doc.new_subdoc(toc_path)
    context = {
        "title1": title1,
        "title2": title2,
        "title3": title3,
        "TOC": toc,
    }
    doc.render(context)
    # Start new section for different column layout
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    master = Composer(doc)
    # Add the main (non-front-matter) content.
    master.append(composer.doc)
    master.save(docx_filepath)
    t1 = time.time()
    logger.debug("Time for converting HTML to Docx: %s", t1 - t0)


def generate_docx_toc(docx_filepath: str) -> str:
    """
    Create subdocument that contains only the code to generate (on
    first open of document) the table of contents.
    """
    toc_path = "{}_toc.docx".format(Path(docx_filepath).with_suffix(""))
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = (
        r'TOC \o "1-2" \h \z \u'  # change 1-2 depending on heading levels you need
    )
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = (
        "Right-click to update field (doing so will insert table of contents)."
    )
    fldChar2.append(fldChar3)
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    document.save(toc_path)
    return toc_path


def html_filepath(
    document_request_key: str, output_dir: str = settings.DOCUMENT_OUTPUT_DIR
) -> str:
    """Given document_request_key, return the HTML output file path."""
    return join(output_dir, "{}.html".format(document_request_key))


def pdf_filepath(
    document_request_key: str, output_dir: str = settings.DOCUMENT_OUTPUT_DIR
) -> str:
    """Given document_request_key, return the PDF output file path."""
    return join(output_dir, "{}.pdf".format(document_request_key))


def epub_filepath(
    document_request_key: str, output_dir: str = settings.DOCUMENT_OUTPUT_DIR
) -> str:
    """Given document_request_key, return the ePub output file path."""
    return join(output_dir, "{}.epub".format(document_request_key))


def docx_filepath(
    document_request_key: str, output_dir: str = settings.DOCUMENT_OUTPUT_DIR
) -> str:
    """Given document_request_key, return the docx output file path."""
    return join(output_dir, "{}.docx".format(document_request_key))


def cover_filepath(
    document_request_key: str, output_dir: str = settings.DOCUMENT_OUTPUT_DIR
) -> str:
    """Given document_request_key, return the HTML cover output file path."""
    return join(output_dir, "{}_cover.html".format(document_request_key))


def select_assembly_layout_kind(
    document_request: DocumentRequest,
    usfm_resource_types: Sequence[str] = settings.USFM_RESOURCE_TYPES,
    language_book_order: AssemblyStrategyEnum = AssemblyStrategyEnum.LANGUAGE_BOOK_ORDER,
    book_language_order: AssemblyStrategyEnum = AssemblyStrategyEnum.BOOK_LANGUAGE_ORDER,
    one_column_compact: AssemblyLayoutEnum = AssemblyLayoutEnum.ONE_COLUMN_COMPACT,
    sl_sr: AssemblyLayoutEnum = AssemblyLayoutEnum.TWO_COLUMN_SCRIPTURE_LEFT_SCRIPTURE_RIGHT,
    sl_sr_compact: AssemblyLayoutEnum = AssemblyLayoutEnum.TWO_COLUMN_SCRIPTURE_LEFT_SCRIPTURE_RIGHT_COMPACT,
    one_column: AssemblyLayoutEnum = AssemblyLayoutEnum.ONE_COLUMN,
) -> AssemblyLayoutEnum:
    """
    Make an intelligent choice of what layout to use given the
    DocumentRequest instance the user has requested. Note that prior to
    this, we've already validated the DocumentRequest instance in the
    DocumentRequest's validator. If we hadn't then we wouldn't be able
    to make the assumptions this function makes.
    """

    # The assembly_layout_kind does not get set by the UI, so if it is set
    # now that means that the request is coming from a client other than the UI.
    # In either case validation of the DocumentRequest instance will have
    # already occurred by this point thus ensuring that the document
    # request's values are valid in which case we can simply return the
    # assembly_layout_kind that was set.
    if (
        document_request.document_request_source == DocumentRequestSourceEnum.TEST
        and document_request.assembly_layout_kind
    ):
        return document_request.assembly_layout_kind
    if (
        document_request.layout_for_print
        and document_request.assembly_strategy_kind == language_book_order
    ):
        return one_column_compact
    elif (
        not document_request.layout_for_print
        and document_request.assembly_strategy_kind == language_book_order
    ):
        return one_column
    elif (
        not document_request.layout_for_print
        and document_request.assembly_strategy_kind == book_language_order
    ):
        # return sl_sr
        return one_column
    elif (
        document_request.layout_for_print
        and document_request.assembly_strategy_kind == book_language_order
    ):
        # return sl_sr_compact
        return one_column_compact
    return one_column


def write_html_content_to_file(
    content: str,
    output_filename: str,
) -> None:
    """
    Write HTML content to file.
    """
    logger.debug("About to write HTML to %s", output_filename)
    # Write the HTML file to disk.
    write_file(
        output_filename,
        content,
    )
    copy_file_to_docker_output_dir(output_filename)


def copy_file_to_docker_output_dir(
    filepath: str,
    document_output_dir: str = settings.DOCUMENT_SERVE_DIR,
) -> None:
    """
    Copy file to docker_container_document_output_dir.
    """
    assert exists(filepath)
    shutil.copy(filepath, document_output_dir)
    logger.debug("About to cp file: %s to directory: %s", filepath, document_output_dir)


def check_content_for_issues(
    content: str,
    # lang_code: str,
    # book_code: str,
    # resource_type: str,
    # github_api_token: str = settings.GITHUB_API_TOKEN,
    # create_github_issue: bool = False,  # Not using this feature yet
) -> str:
    """
    Check for defects and notify support via logs of potential source
    content issues. Also modify the content to include a message for
    the end user to inform them that there is a problem with the
    underlying source USFM and that the translators need to fix it.
    This will help them understand why they have missing content.
    """
    logger.info(
        "Checking USFM content for issues before creating requested document..."
    )
    verses_html = re.findall(
        r'<div .*?class="verse".*?>(.*?)</div>', content, re.DOTALL
    )
    logger.debug("verses_html: %s", verses_html)
    if not verses_html:
        logger.info("No verses found in HTML")
    verses_combined = "".join(verses_html)
    if not verses_html:
        logger.info(
            "About to modify content to include message notifying user of problem with USFM source text format..."
        )
        updated_content = "NOTE: There are issues with the requested underlying scripture USFM text that make it unusable by this system until translators fix the issue for the language(s), book(s), and resource(s) combination you have requested."
        logger.debug(
            "Due to potential issues with the source content, here is the HTML content for you to inspect: %s",
            content,
        )
        # if create_github_issue:
        #     # Automatically record an issue using github API
        #     document_request_info = (lang_code, book_code, resource_type)
        #     create_github_issue_command = (
        #         "curl -L -X POST -H"
        #         "Accept: application/vnd.github+json"
        #         "-H Authorization: Bearer {}"
        #         "-H X-GitHub-Api-Version: {}"
        #         "https://api.github.com/repos/WycliffeAssociates/DOC/issues"
        #         '-d {{"title":"Found USFM source content bug","body":"USFM for one of {} has an issue that makes it unparsable."}}'
        #     ).format(
        #         github_api_token,
        #         "2022-11-28",
        #         document_request_info,
        #     )
        #     subprocess.call(create_github_issue_command, shell=True)
        return updated_content
    return content


# @worker.app.task(
#     autoretry_for=(Exception,),
#     retry_backoff=True,
#     retry_kwargs={"max_retries": 3},
# )
@worker.app.task
def generate_document(
    document_request_json: Json[Any], output_dir: str = settings.DOCUMENT_OUTPUT_DIR
) -> Json[Any]:
    """
    This is the main entry point for this module for non-docx generation.
    >>> from document.domain import document_generator
    >>> document_request_json = '{"email_address":null,"assembly_strategy_kind":"lbo","assembly_layout_kind":"1c","layout_for_print":false,"resource_requests":[{"lang_code":"es-419","resource_type":"ulb","book_code":"mat"}],"generate_pdf":true,"generate_epub":false,"generate_docx":false,"chunk_size":"chapter","limit_words":false,"include_tn_book_intros":false,"document_request_source":"ui"}'
    >>> document_generator.generate_document(document_request_json)
    """
    logger.debug(
        "document_request_json: %s",
        document_request_json,
    )
    current_task.update_state(state="Receiving request")
    document_request = DocumentRequest.parse_raw(document_request_json)
    document_request.assembly_layout_kind = select_assembly_layout_kind(
        document_request
    )
    # Generate the document request key that identifies this and
    # identical document requests.
    document_request_key_ = document_request_key(
        document_request.resource_requests,
        document_request.assembly_strategy_kind,
        document_request.assembly_layout_kind,
        document_request.chunk_size,
        document_request.limit_words,
    )
    html_filepath_ = html_filepath(document_request_key_)
    pdf_filepath_ = pdf_filepath(document_request_key_)
    epub_filepath_ = epub_filepath(document_request_key_)
    if file_needs_update(html_filepath_):
        # Update the state of the worker process. This is used by the
        # UI to report status.
        current_task.update_state(state="Locating assets")
        # HTML didn't exist in cache so go ahead and start by getting the
        # resource lookup DTOs for each resource request in the document
        # request.
        resource_lookup_dtos = []
        for resource_request in document_request.resource_requests:
            resource_lookup_dto = resource_lookup.resource_lookup_dto(
                resource_request.lang_code,
                resource_request.resource_type,
                resource_request.book_code,
            )
            if resource_lookup_dto:
                resource_lookup_dtos.append(resource_lookup_dto)
        # Determine which resource URLs were actually found.
        found_resource_lookup_dtos = [
            resource_lookup_dto
            for resource_lookup_dto in resource_lookup_dtos
            if resource_lookup_dto.url is not None
        ]
        # if not found_resource_lookup_dtos:
        #     raise exceptions.ResourceAssetFileNotFoundError(
        #         message="No supported resource assets were found"
        #     )
        current_task.update_state(state="Provisioning asset files")
        t0 = time.time()
        resource_dirs = [
            resource_lookup.provision_asset_files(dto)
            for dto in found_resource_lookup_dtos
        ]
        t1 = time.time()
        logger.debug(
            "Time to provision asset files (acquire and write to disk): %s", t1 - t0
        )
        current_task.update_state(state="Parsing asset files")
        # Initialize found resources from their provisioned assets.
        t0 = time.time()
        usfm_books, tn_books, tq_books, tw_books, bc_books = parsing.books(
            found_resource_lookup_dtos,
            resource_dirs,
            document_request.resource_requests,
            document_request.layout_for_print,
        )
        t1 = time.time()
        logger.debug("Time to parse all resource content: %s", t1 - t0)
        current_task.update_state(state="Assembling content")
        content = assemble_content(
            document_request_key_,
            document_request,
            usfm_books,
            tn_books,
            tq_books,
            tw_books,
            bc_books,
            found_resource_lookup_dtos,
        )
        if usfm_books:
            content = check_content_for_issues(content)
        content = create_title_page_and_wrap_in_template(
            content, document_request, found_resource_lookup_dtos
        )
        write_html_content_to_file(
            content,
            html_filepath_,
        )
    else:
        logger.debug("Cache hit for %s", html_filepath_)
    # Immediately return pre-built PDF if the document has previously been
    # generated and is fresh enough.
    if document_request.generate_pdf and file_needs_update(pdf_filepath_):
        current_task.update_state(state="Converting to PDF")
        convert_html_to_pdf(
            html_filepath_,
            pdf_filepath_,
            document_request_key_,
        )
        copy_file_to_docker_output_dir(pdf_filepath_)
        if should_send_email(document_request.email_address):
            attachments = [
                Attachment(filepath=pdf_filepath_, mime_type=("application", "pdf"))
            ]
            current_task.update_state(state="Sending email")
            send_email_with_attachment(
                document_request.email_address,
                attachments,
                document_request_key_,
            )
    if document_request.generate_epub and file_needs_update(epub_filepath_):
        current_task.update_state(state="Converting to ePub")
        convert_html_to_epub(
            html_filepath_,
            epub_filepath_,
            document_request_key_,
        )
        copy_file_to_docker_output_dir(epub_filepath_)
        if should_send_email(document_request.email_address):
            attachments = [
                Attachment(
                    filepath=epub_filepath_, mime_type=("application", "epub+zip")
                )
            ]
            current_task.update_state(state="Sending email")
            send_email_with_attachment(
                document_request.email_address,
                attachments,
                document_request_key_,
            )
    return document_request_key_


@worker.app.task
def generate_docx_document(
    document_request_json: Json[Any],
    output_dir: str = settings.DOCUMENT_OUTPUT_DIR,
) -> Json[Any]:
    """
    This is the alternative entry point for Docx document creation only.
    """
    document_request = DocumentRequest.parse_raw(document_request_json)
    logger.debug(
        "document_request: %s",
        document_request,
    )
    document_request.assembly_layout_kind = select_assembly_layout_kind(
        document_request
    )
    # Generate the document request key that identifies this and
    # identical document requests.
    document_request_key_ = document_request_key(
        document_request.resource_requests,
        document_request.assembly_strategy_kind,
        document_request.assembly_layout_kind,
        document_request.chunk_size,
        document_request.limit_words,
    )
    html_filepath_ = html_filepath(document_request_key_)
    docx_filepath_ = docx_filepath(document_request_key_)
    if document_request.generate_docx and file_needs_update(docx_filepath_):
        # Update the state of the worker process. This is used by the
        # UI to report status.
        current_task.update_state(state="Locating assets")
        # Docx didn't exist in cache so go ahead and start by getting the
        # resource lookup DTOs for each resource request in the document
        # request.
        resource_lookup_dtos = []
        for resource_request in document_request.resource_requests:
            resource_lookup_dto = resource_lookup.resource_lookup_dto(
                resource_request.lang_code,
                resource_request.resource_type,
                resource_request.book_code,
            )
            if resource_lookup_dto:
                resource_lookup_dtos.append(resource_lookup_dto)
        # Determine which resource URLs were actually found.
        found_resource_lookup_dtos = [
            resource_lookup_dto
            for resource_lookup_dto in resource_lookup_dtos
            if resource_lookup_dto.url is not None
        ]
        current_task.update_state(state="Provisioning asset files")
        t0 = time.time()
        resource_dirs = [
            resource_lookup.provision_asset_files(dto)
            for dto in found_resource_lookup_dtos
        ]
        t1 = time.time()
        logger.debug(
            "Time to provision asset files (acquire and write to disk): %s", t1 - t0
        )
        current_task.update_state(state="Parsing asset files")
        # Initialize found resources from their provisioned assets.
        t0 = time.time()
        usfm_books, tn_books, tq_books, tw_books, bc_books = parsing.books(
            found_resource_lookup_dtos,
            resource_dirs,
            document_request.resource_requests,
            document_request.layout_for_print,
        )
        t1 = time.time()
        logger.debug("Time to parse all resource content: %s", t1 - t0)
        current_task.update_state(state="Assembling content")
        composer = assemble_docx_content(
            document_request_key_,
            document_request,
            usfm_books,
            tn_books,
            tq_books,
            tw_books,
            bc_books,
        )
        # TODO At this point, like in generate_document, we should check the
        # underlying HTML content to see if it contains verses and display a
        # message in the document to the end user if it does not. Would just
        # have to figure out how to get at the underlying HTML contained in the
        # composer instance.
        # Construct sensical phrases to display for title1 and title2 on first
        # page of Word document.
        title1, title2 = get_languages_title_page_strings(found_resource_lookup_dtos)
        current_task.update_state(state="Converting to Docx")
        convert_html_to_docx(
            html_filepath_,
            docx_filepath_,
            composer,
            document_request.layout_for_print,
            title1,
            title2,
        )
        copy_file_to_docker_output_dir(docx_filepath_)
        if should_send_email(document_request.email_address):
            attachments = [
                Attachment(
                    filepath=docx_filepath_,
                    mime_type=(
                        "application",
                        "vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ]
            current_task.update_state(state="Sending email")
            send_email_with_attachment(
                document_request.email_address,
                attachments,
                document_request_key_,
            )
    else:
        logger.debug("Cache hit for %s", docx_filepath_)
    return document_request_key_


def get_languages_title_page_strings(
    resource_lookup_dtos: Sequence[ResourceLookupDto],
) -> tuple[str, str]:
    """
    Return a list of tuples with the following form:
    [(lang_name, [book_code1, book_code2, ...]), ...]
    E.g.,
    [("English", ["mat", "mrk"]), ("français (French)", ["mat", "mrk"])]
    >>> from document.domain import document_generator, model
    >>> # resource_lookup_dtos=[model.ResourceLookupDto(lang_code="en", lang_name="English", resource_type="ulb-wa", resource_type_name="Scripture", book_code="mat", source="usfm"), model.ResourceLookupDto(lang_code="fr", lang_name="français (French)", resource_type="ulb", resource_type_name="Translation Notes", book_code="mat", source="usfm")]
    >>> # resource_lookup_dtos
    >>> #document_generator.get_languages_title_page_strings(resource_lookup_dtos)
    """
    lang_codes = list(
        {resource_lookup_dto.lang_code for resource_lookup_dto in resource_lookup_dtos}
    )
    language0_resource_lookup_dtos = [
        resource_lookup_dto
        for resource_lookup_dto in resource_lookup_dtos
        if resource_lookup_dto.lang_code == lang_codes[0]
    ]
    lang0_book_names = []
    lang0_resource_type_names = []
    for lang0_dto in language0_resource_lookup_dtos:
        book_name = BOOK_NAMES[lang0_dto.book_code]
        if book_name and book_name not in lang0_book_names:
            lang0_book_names.append(book_name)
        if lang0_dto.resource_type_name not in lang0_resource_type_names:
            lang0_resource_type_names.append(lang0_dto.resource_type_name)
    lang0_title = "{}: {} for {}".format(
        language0_resource_lookup_dtos[0].lang_name,
        ", ".join(sorted(lang0_resource_type_names)),
        ", ".join(sorted(lang0_book_names)),
    )
    lang1_title = ""
    if len(lang_codes) > 1:
        language1_resource_lookup_dtos = [
            resource_lookup_dto
            for resource_lookup_dto in resource_lookup_dtos
            if resource_lookup_dto.lang_code == lang_codes[1]
        ]
        lang1_book_names = []
        lang1_resource_type_names = []
        for lang1_dto in language1_resource_lookup_dtos:
            book_name = BOOK_NAMES[lang1_dto.book_code]
            if book_name and book_name not in lang1_book_names:
                lang1_book_names.append(book_name)
            if lang1_dto.resource_type_name not in lang1_resource_type_names:
                lang1_resource_type_names.append(lang1_dto.resource_type_name)
        lang1_title = "{}: {} for {}".format(
            language1_resource_lookup_dtos[0].lang_name,
            ", ".join(sorted(lang1_resource_type_names)),
            ", ".join(sorted(lang1_book_names)),
        )
    return lang0_title, lang1_title


if __name__ == "__main__":
    # To run the doctests in the this module, in the root of the project do:
    # FROM_EMAIL_ADDRESS=... python backend/document/domain/resource_lookup.py
    # See https://docs.python.org/3/library/doctest.html
    # for more details.
    import doctest

    doctest.testmod()
