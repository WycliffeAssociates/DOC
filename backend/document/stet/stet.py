import re
from pathlib import Path
from typing import Mapping, Sequence

import jinja2
import mistune
from celery import current_task
from document.config import settings
from document.domain.bible_books import BOOK_NAMES
from document.domain.model import USFMBook, USFMChapter
from document.domain.parsing import usfm_book_content
from document.domain.resource_lookup import (
    provision_asset_files,
    resource_lookup_dto,
    resource_types,
)
from document.stet.model import VerseEntry, VerseReferenceDto, WordEntry, WordEntryDto
from document.stet.util import is_valid_int
from document.utils.file_utils import template_path
from docx import Document  # type: ignore
from htmldocx import HtmlToDocx  # type: ignore

logger = settings.logger(__name__)


def lookup_verse_text(usfm_book: USFMBook, chapter_num: int, verse_ref: str) -> str:
    if chapter_num in usfm_book.chapters:
        chapter = usfm_book.chapters[chapter_num]
        if chapter.verses:
            verse = chapter.verses[verse_ref] if verse_ref in chapter.verses else ""
            logger.debug(
                "book_code: %s, chapter_num: %s, verse_num: %s, verse: %s",
                usfm_book.book_code,
                chapter_num,
                verse_ref,
                verse,
            )
            return verse
        return ""
    return ""


def split_chapter_into_verses(chapter: USFMChapter) -> dict[str, str]:
    # Sample HTML content with multiple verse elements
    # html_content = '''
    # <span class="verse">
    # <sup class="versemarker">19</sup>
    # For through the law I died to the law, so that I might live for God. I have been crucified with Christ.
    # <sup id="footnote-caller-1" class="caller"><a href="#footnote-target-1">1</a></sup>
    # <div class="sectionhead-5"></div>
    # </span>

    # <span class="verse">
    # <sup class="versemarker">20</sup>
    # I have been crucified with Christ and I no longer live, but Christ lives in me. The life I now live in the body, I live by faith in the Son of God, who loved me and gave himself for me.
    # <sup id="footnote-caller-2" class="caller"><a href="#footnote-target-2">2</a></sup>
    # <div class="sectionhead-5"></div>
    # </span>
    # '''

    verse_dict = {}
    # Find all verse spans
    verse_spans = re.findall(
        r'<span class="verse">(.*?)</span>', chapter.content, re.DOTALL
    )
    for verse_span in verse_spans:
        # Extract the verse number from the versemarker
        verse_number = re.search(r'<sup class="versemarker">(\d+)</sup>', verse_span)
        if verse_number:
            verse_number_ = verse_number.group(1)
            # Remove all <sup> and <div> tags and their content from the verse text
            verse_text = re.sub(r"<sup.*?>.*?</sup>", "", verse_span)
            verse_text = re.sub(r"<div.*?>.*?</div>", "", verse_text)
            # Remove the remaining HTML tags and strip extra spaces
            verse_text = re.sub(r"<.*?>", "", verse_text).strip()
            logger.debug("verse_number: %s, verse_text: %s", verse_number_, verse_text)
            # Add to the dictionary with verse number as the key and verse text as the value
            verse_dict[verse_number_] = verse_text
    return verse_dict


def get_word_entry_dtos(
    lang0_code: str,
    lang1_code: str,
    book_names: dict[str, str] = BOOK_NAMES,
    stet_dir: str = settings.STET_DIR,
) -> tuple[list[WordEntryDto], list[str]]:
    # Build data from source doc
    word_entry_dtos: list[WordEntryDto] = []
    book_codes_: list[str] = []
    doc = Document(f"{stet_dir}/stet_{lang0_code}.docx")
    for table in doc.tables:
        for row in table.rows:
            # Create entry item
            word_entry_dto = WordEntryDto()
            # Extract data from word field
            match = re.match(r"(.*)(\n)?(.*)?", row.cells[0].text)
            # TODO Maybe shouldn't raise an exception here
            if not match:
                raise ValueError(f"Couldn't parse word def: {row.cells[0].text}")
            word = match.group(1)
            word_entry_dto.word = word
            raw_strongs = match.group(3)
            word_entry_dto.strongs_numbers = raw_strongs.strip()
            definition = ""
            previous_paragraph_style_name = ""
            for paragraph in row.cells[1].paragraphs:
                if previous_paragraph_style_name not in (paragraph.style.name, ""):
                    definition += "\n"
                if paragraph.style.name == "List Paragraph":
                    definition += f"- {paragraph.text.strip()}\n"
                else:
                    definition += f"{paragraph.text.strip()}\n"
                previous_paragraph_style_name = paragraph.style.name
            word_entry_dto.definition = definition
            # process verse list
            for reference in row.cells[2].text.split("\n"):
                reference_ = reference.strip()
                match = re.match(r"^(.*) (\d+):([0-9,\- ]+)\s?(\(.*\))?$", reference_)
                if not match:
                    logger.warning("Couldn't parse %s", reference_)
                    continue
                if match:
                    # Extract references
                    book_name = match.group(1)
                    book_codes = [
                        book_code
                        for book_code, book_name_ in book_names.items()
                        if book_name_ == book_name
                    ]
                    book_code = book_codes[0] if book_codes else None
                    if book_code:
                        book_codes_.append(book_code)
                    chapter_num = int(match.group(2))
                    verses = match.group(3)
                    comment = match.group(4)
                    # Clean up comments
                    # if comment:
                    #     # Replace asterisks with Markdown-compatible asterisks
                    #     comment = re.sub(r"\*", r"\\*", comment)
                    # Write row to table
                    if comment:
                        source_reference = (
                            f"{book_name} {chapter_num}:{verses}{comment}"
                        )
                    else:
                        source_reference = f"{book_name} {chapter_num}:{verses}"
                    target_reference = f"{book_name} {chapter_num}:{verses}"
                    logger.debug(
                        "book_name: %s, chapter_num: %s, verse_num(s): %s, comment: %s",
                        book_name,
                        chapter_num,
                        verses,
                        comment,
                    )
                    # Break apart source tool verses
                    verse_refs: list[str] = verses.split(",")
                    valid_verse_refs: list[str] = []
                    for verse_ref in verse_refs:
                        if is_valid_int(verse_ref):
                            valid_verse_refs.append(str(verse_ref))
                            continue
                        match = re.match(r"(\d+)-(\d+)", verse_ref)
                        if match:
                            start_verse = int(match.group(1))
                            end_verse = int(match.group(2))
                            verse_num = start_verse
                            while verse_num <= end_verse:
                                valid_verse_refs.append(str(verse_num))
                                verse_num += 1
                            continue
                        logger.warning("Couldn't parse verse ref: %s", verse_ref)
                    verse_reference_dto = VerseReferenceDto(
                        lang0_code=lang0_code,
                        lang1_code=lang1_code,
                        book_code=book_code,
                        book_name=book_name,
                        chapter_num=chapter_num,
                        source_reference=source_reference,
                        target_reference=target_reference,
                        verse_refs=valid_verse_refs,
                    )
                    word_entry_dto.verse_ref_dtos.append(verse_reference_dto)
            word_entry_dtos.append(word_entry_dto)
    return word_entry_dtos, list(set(book_codes_))


def generate_docx_document(
    lang0_code: str,
    lang1_code: str,
    document_request_key_: str,
    docx_filepath_: str,
    working_dir: str = settings.WORKING_DIR,
    output_dir: str = settings.DOCUMENT_OUTPUT_DIR,
    usfm_resource_types: Sequence[str] = settings.USFM_RESOURCE_TYPES,
    resource_type_codes_and_names: Mapping[
        str, str
    ] = settings.RESOURCE_TYPE_CODES_AND_NAMES,
) -> str:
    """
    Generate the scriptural terms evaluation document.

    >>> from document.stet import generate_markdown_document
    >>> generate_docx_document()
    """
    word_entries: list[WordEntry] = []
    word_entry_dtos, book_codes = get_word_entry_dtos(lang0_code, lang1_code)
    lang0_resource_types = resource_types(lang0_code, ",".join(book_codes))
    lang0_resource_types_ = [
        lang0_resource_type_tuple[0]
        for lang0_resource_type_tuple in lang0_resource_types
    ]
    lang1_resource_types = resource_types(lang1_code, ",".join(book_codes))
    lang1_resource_types_ = [
        lang1_resource_type_tuple[0]
        for lang1_resource_type_tuple in lang1_resource_types
    ]
    lang0_usfm_resource_types = [
        resource_type_
        for resource_type_ in lang0_resource_types_
        if resource_type_ in usfm_resource_types
    ]
    lang1_usfm_resource_types = [
        resource_type_
        for resource_type_ in lang1_resource_types_
        if resource_type_ in usfm_resource_types
    ]
    lang0_ulb_usfm_resource_types = [
        usfm_resource_type_
        for usfm_resource_type_ in lang0_usfm_resource_types
        if "ulb" in usfm_resource_type_
    ]
    lang1_ulb_usfm_resource_types = [
        usfm_resource_type_
        for usfm_resource_type_ in lang1_usfm_resource_types
        if "ulb" in usfm_resource_type_
    ]
    source_usfm_books = []
    target_usfm_books = []
    lang0_usfm_resource_type = ""
    lang1_usfm_resource_type = ""
    if lang0_ulb_usfm_resource_types:
        lang0_usfm_resource_type = lang0_ulb_usfm_resource_types[0]
    elif lang0_usfm_resource_types:
        lang0_usfm_resource_type = lang0_usfm_resource_types[0]
    if lang0_usfm_resource_type:
        for book_code in book_codes:
            # Update the state of the worker process. This is used by the
            # UI to report status.
            logger.debug(
                "About to create dto for lang0_code: %s, book_code: %s, lang0_usfm_resource_type: %s",
                lang0_code,
                book_code,
                lang0_usfm_resource_type,
            )
            current_task.update_state(state="Locating assets")
            lang0_resource_lookup_dto_ = resource_lookup_dto(
                lang0_code, lang0_usfm_resource_type, book_code
            )
            if lang0_resource_lookup_dto_ and lang0_resource_lookup_dto_.url:
                current_task.update_state(state="Provisioning asset files")
                lang0_resource_dir = provision_asset_files(lang0_resource_lookup_dto_)
                current_task.update_state(state="Parsing asset files")
                source_usfm_book = usfm_book_content(
                    lang0_resource_lookup_dto_,
                    lang0_resource_dir,
                )
                for chapter_num_, chapter_ in source_usfm_book.chapters.items():
                    source_usfm_book.chapters[
                        chapter_num_
                    ].verses = split_chapter_into_verses(chapter_)
                source_usfm_books.append(source_usfm_book)
    if lang1_ulb_usfm_resource_types:
        lang1_usfm_resource_type = lang1_ulb_usfm_resource_types[0]
    elif lang1_usfm_resource_types:
        lang1_usfm_resource_type = lang1_usfm_resource_types[0]
    if lang1_usfm_resource_type:
        for book_code in book_codes:
            lang1_resource_lookup_dto_ = resource_lookup_dto(
                lang1_code, lang1_usfm_resource_type, book_code
            )
            if lang1_resource_lookup_dto_ and lang1_resource_lookup_dto_.url:
                lang1_resource_dir = provision_asset_files(lang1_resource_lookup_dto_)
                target_usfm_book = usfm_book_content(
                    lang1_resource_lookup_dto_,
                    lang1_resource_dir,
                )
                for chapter_num_, chapter_ in target_usfm_book.chapters.items():
                    target_usfm_book.chapters[
                        chapter_num_
                    ].verses = split_chapter_into_verses(chapter_)
                target_usfm_books.append(target_usfm_book)
    current_task.update_state(state="Assembling content")
    for word_entry_dto in word_entry_dtos:
        source_verse_text = ""
        target_verse_text = ""
        word_entry = WordEntry()
        word_entry.word = word_entry_dto.word
        word_entry.strongs_numbers = word_entry_dto.strongs_numbers
        word_entry.definition = mistune.markdown(word_entry_dto.definition)
        for verse_ref_dto in word_entry_dto.verse_ref_dtos:
            source_selected_usfm_books = [
                usfm_book_
                for usfm_book_ in source_usfm_books
                if usfm_book_.lang_code == lang0_code
                and usfm_book_.book_code == verse_ref_dto.book_code
                and usfm_book_.resource_type_name
                == resource_type_codes_and_names[lang0_usfm_resource_type]
            ]
            target_selected_usfm_books = [
                usfm_book_
                for usfm_book_ in target_usfm_books
                if usfm_book_.lang_code == lang1_code
                and usfm_book_.book_code == verse_ref_dto.book_code
                and usfm_book_.resource_type_name
                == resource_type_codes_and_names[lang1_usfm_resource_type]
            ]
            if source_selected_usfm_books:
                source_selected_usfm_book = source_selected_usfm_books[0]
                for verse_ref in verse_ref_dto.verse_refs:
                    source_verse_text = lookup_verse_text(
                        source_selected_usfm_book, verse_ref_dto.chapter_num, verse_ref
                    )
            if target_selected_usfm_books:
                target_selected_usfm_book = target_selected_usfm_books[0]
                for verse_ref in verse_ref_dto.verse_refs:
                    target_verse_text = lookup_verse_text(
                        target_selected_usfm_book, verse_ref_dto.chapter_num, verse_ref
                    )
            word_entry.verses.append(
                VerseEntry(
                    source_reference=verse_ref_dto.source_reference,
                    source_text=source_verse_text,
                    target_reference=verse_ref_dto.target_reference,
                    target_text=target_verse_text,
                )
            )
        word_entries.append(word_entry)
    # Build output doc
    # Create markdown file that you can run pandoc on
    # template = Path(template_path("stet")).read_text(encoding="utf-8")
    # # markdown_ = chevron.render(template=template, data={"words": word_entries})
    # markdown_ = chevron.render(template=template, data=word_entries)  # type: ignore
    # filepath_ = f"{working_dir}/{lang0_code}_{lang1_code}_stet.md"
    # with open(filepath_, "w", encoding="utf-8") as outfile:
    #     outfile.write(markdown_)
    # return filepath_
    # Create HTML file and then convert it to Docx with library
    current_task.update_state(state="Converting to Docx")
    template = Path(template_path("stet_html")).read_text(encoding="utf-8")
    # Hydrate and render the template
    # assert exists(template_html)
    # with open(template_html, "r") as filepath:
    #     template = filepath.read()
    env = jinja2.Environment(autoescape=True).from_string(template)
    full_html = env.render(data=word_entries)
    # Remove any NULL bytes and all control characters
    cleaned_full_html = re.sub(r"[\x00-\x1F]+", "", full_html)
    # logger.debug("full_html: %s", full_html)
    # filepath_ = f"{working_dir}/{lang0_code}_{lang1_code}_stet.html"
    filepath_ = f"{working_dir}/{document_request_key_}.html"
    with open(filepath_, "w", encoding="utf-8") as outfile2:
        outfile2.write(cleaned_full_html)

    html_to_docx = HtmlToDocx()
    # docx_filepath = f"{Path(filepath_).stem}.docx"
    html_to_docx.parse_html_file(filepath_, f"{output_dir}/{Path(docx_filepath_).stem}")
    return docx_filepath_
    # # doc = html_to_docx.parse_html_string(html)
    # # logger.debug("doc: %s", doc)
