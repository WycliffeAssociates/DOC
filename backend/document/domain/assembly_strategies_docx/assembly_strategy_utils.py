"""
Utility functions used by assembly_strategies.
"""

from document.config import settings
from docx import Document  # type: ignore
from docx.enum.section import WD_SECTION  # type: ignore
from docx.enum.text import WD_BREAK  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml.shared import OxmlElement  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from htmldocx import HtmlToDocx  # type: ignore

logger = settings.logger(__name__)

H1, H2, H3, H4, H5, H6 = "h1", "h2", "h3", "h4", "h5", "h6"


OXML_LANGUAGE_LIST: list[str] = [
    "ar-SA",
    "bg-BG",
    "zh-CN",
    "zh-TW",
    "hr-HR",
    "cs-CZ",
    "da-DK",
    "nl-NL",
    "en-US",
    "et-EE",
    "fi-FI",
    "fr-FR",
    "de-DE",
    "el-GR",
    "he-IL",
    "hi-IN",
    "hu-HU",
    "id-ID",
    "it-IT",
    "ja-JP",
    "kk-KZ",
    "ko-KR",
    "lv-LV",
    "lt-LT",
    "ms-MY",
    "nb-NO",
    "pl-PL",
    "pt-BR",
    "pt-PT",
    "ro-RO",
    "ru-RU",
    "sr-latn-RS",
    "sk-SK",
    "sl-SI",
    "es-ES",
    "sv-SE",
    "th-TH",
    "tr-TR",
    "uk-UA",
    "vi-VN",
]
OXML_LANGUAGE_LIST_LOWERCASE: list[str] = [
    language.lower() for language in OXML_LANGUAGE_LIST
]
OXML_LANGUAGE_LIST_LOWERCASE_SPLIT: list[str] = [
    language for language in OXML_LANGUAGE_LIST_LOWERCASE if "-" in language
]


def add_hr(paragraph: Paragraph) -> None:
    """Add a horizontal line at the end of the given paragraph."""
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.insert_element_before(
        pBdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def create_docx_subdoc(
    content: str,
    lang_code: str,
    is_rtl: bool = False,
    add_hr_p: bool = True,
    oxml_language_list_lowercase: list[str] = OXML_LANGUAGE_LIST_LOWERCASE,
    oxml_language_list_lowercase_split: list[str] = OXML_LANGUAGE_LIST_LOWERCASE_SPLIT,
) -> Document:
    """
    Create and return a Document instance from the content parameter.
    """
    html_to_docx = HtmlToDocx()
    subdoc = html_to_docx.parse_html_string(content)
    if is_rtl:
        # Setting each run to be RTL language direction
        for p in subdoc.paragraphs:
            for run in p.runs:
                run.font.rtl = True
    if subdoc.paragraphs:
        p = subdoc.paragraphs[-1]
        # Set the language for this paragraph for the sake of the Word
        # spellchecker.
        p_run = p.add_run()
        # if is_rtl:
        #     p_run.font.rtl = True
        p_rpr = p_run.element.get_or_add_rPr()
        p_run_lang = OxmlElement("w:lang")
        oxml_language_list_lowercase_split_values = [
            language.lower().split("-")[0]
            for language in oxml_language_list_lowercase_split
        ]
        # The code below works well for spell checking, but there is some
        # action required by the user to configure Word. The user must add
        # the languages that their document uses under File > Options > Language
        # then use the 'Add additional editing languages' and then enable
        # the languages (it will say 'Not enabled'), then restart Word.
        if lang_code in oxml_language_list_lowercase:
            # Set language for spell and grammar check for this run.
            p_run_lang.set(qn("w:val"), lang_code)
            p_run_lang.set(qn("w:eastAsia"), lang_code)
            # bidi is short for bidirectionality text
            p_run_lang.set(qn("w:bidi"), lang_code)

            # # Set font for this run.
            # p_run.font.name = "Noto Sans Regular"
            # r = p_run._element
            # r.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Regular")
        elif lang_code in oxml_language_list_lowercase_split_values:
            # Set language for spell and grammar check for this run.
            updated_lang_code = "{}-{}".format(lang_code, lang_code.upper())
            p_run_lang.set(qn("w:val"), updated_lang_code)
            p_run_lang.set(qn("w:eastAsia"), updated_lang_code)
            # bidi is short for bidirectionality text
            p_run_lang.set(qn("w:bidi"), updated_lang_code)

            # # Set font for this run.
            # p_run.font.name = "Noto Sans Regular"
            # r = p_run._element
            # r.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Regular")
        else:
            # Set language for spell and grammar check for this run.
            # Just set to English since language isn't a language that
            # Word supports (need to research what Word fully supports in
            # more depth to be sure).
            p_run_lang.set(qn("w:val"), "en-US")
            p_run_lang.set(qn("w:eastAsia"), "en-US")
            p_run_lang.set(qn("w:bidi"), "en-US")

        p_rpr.append(p_run_lang)
        # Add a horizontal ruler at the end of the paragraph if requested.
        if add_hr_p:
            add_hr(p)
    return subdoc


def add_one_column_section(doc: Document) -> None:
    """
    Add new section having 1 column to contain next content in Docx instance.
    """
    # Get ready for one column again (this matters the 2nd to Nth times in the loop).
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type

    # Set to one column layout for subdocument to be added next.
    sectPr = new_section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")


def add_two_column_section(doc: Document) -> None:
    """
    Add new section having 2 column to contain next content in Docx instance.
    """
    # Start new section for different (two) column layout.
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    sectPr = new_section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "10")  # Set space between columns to 10 points ->0.01"


def add_page_break(doc: Document) -> None:
    """Add page break."""
    # Add page break at end of chapter content
    p = doc.add_paragraph("")
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
