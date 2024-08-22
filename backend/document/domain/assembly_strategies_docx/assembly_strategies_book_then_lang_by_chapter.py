from typing import Mapping, Sequence

from document.domain.bible_books import BOOK_NAMES, BOOK_CHAPTERS
from document.config import settings
from document.domain.assembly_strategies.assembly_strategy_utils import (
    bc_book_intro,
    chapter_commentary,
    chapter_intro,
    tn_chapter_verses,
    tq_chapter_verses,
)
from document.domain.assembly_strategies_docx.assembly_strategy_utils import (
    add_one_column_section,
    add_two_column_section,
    add_page_break,
    create_docx_subdoc,
)
from document.domain.model import (
    AssemblyLayoutEnum,
    BCBook,
    ChunkSizeEnum,
    LangDirEnum,
    TNBook,
    TQBook,
    TWBook,
    USFMBook,
)

from docx import Document  # type: ignore
from docxcompose.composer import Composer  # type: ignore

logger = settings.logger(__name__)


def assemble_content_by_book_then_lang(
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
    assembly_layout_kind: AssemblyLayoutEnum,
    chunk_size: ChunkSizeEnum,
    book_names: Mapping[str, str] = BOOK_NAMES,
) -> Composer:
    """
    Assemble by book then by language in alphabetic order before
    delegating more atomic ordering/interleaving to an assembly
    sub-strategy.
    """
    # Sort the books in canonical order so that groupby does what we want.
    book_id_map = dict((id, pos) for pos, id in enumerate(BOOK_NAMES.keys()))
    most_book_codes = max(
        [
            [usfm_book.book_code for usfm_book in usfm_books],
            [tn_book.book_code for tn_book in tn_books],
            [tq_book.book_code for tq_book in tq_books],
            [tw_book.book_code for tw_book in tw_books],
            [bc_book.book_code for bc_book in bc_books],
        ],
        key=lambda x: len(x),
    )
    for book_code in sorted(
        most_book_codes,
        key=lambda book_code: book_id_map[book_code],
    ):
        selected_usfm_books = [
            usfm_book for usfm_book in usfm_books if usfm_book.book_code == book_code
        ]
        selected_tn_books = [
            tn_book for tn_book in tn_books if tn_book.book_code == book_code
        ]
        selected_tq_books = [
            tq_book for tq_book in tq_books if tq_book.book_code == book_code
        ]
        selected_tw_books = [
            tw_book for tw_book in tw_books if tw_book.book_code == book_code
        ]
        selected_bc_books = [
            bc_book for bc_book in bc_books if bc_book.book_code == book_code
        ]
        if selected_usfm_books:
            composer = assemble_usfm_by_chapter(
                usfm_books,
                tn_books,
                tq_books,
                tw_books,
                bc_books,
            )
            return composer
        elif not selected_usfm_books and selected_tn_books:
            composer = assemble_tn_by_chapter(
                usfm_books,
                tn_books,
                tq_books,
                tw_books,
                bc_books,
            )
            return composer
        elif not selected_usfm_books and not selected_tn_books and selected_tq_books:
            composer = assemble_tq_by_chapter(
                usfm_books,
                tn_books,
                tq_books,
                tw_books,
                bc_books,
            )
            return composer
        elif (
            not selected_usfm_books
            and not selected_tn_books
            and not selected_tq_books
            and (selected_tw_books or selected_bc_books)
        ):
            composer = assemble_tw_by_chapter(
                usfm_books,
                tn_books,
                tq_books,
                tw_books,
                bc_books,
            )
            return composer


def assemble_usfm_by_chapter(
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
    book_chapters: Mapping[str, int] = BOOK_CHAPTERS,
) -> Composer:
    """
    Construct the Docx wherein at least one USFM resource exists, one column
    layout.
    """

    def sort_key(resource: USFMBook) -> str:
        return resource.lang_code

    def tn_sort_key(resource: TNBook) -> str:
        return resource.lang_code

    def tq_sort_key(resource: TQBook) -> str:
        return resource.lang_code

    def bc_sort_key(resource: BCBook) -> str:
        return resource.lang_code

    usfm_books = sorted(usfm_books, key=sort_key)
    tn_books = sorted(tn_books, key=tn_sort_key)
    tq_books = sorted(tq_books, key=tq_sort_key)
    bc_books = sorted(bc_books, key=bc_sort_key)
    doc = Document()
    composer = Composer(doc)

    # Content team doesn't want TN book intros: https://github.com/WycliffeAssociates/DOC/issues/121
    # # Add book intros for each tn_book_content_unit
    # for tn_book in tn_books:
    #     if tn_book.book_intro:
    #         book_intro_ = tn_book.book_intro
    #         book_intro_adj = adjust_book_intro_headings(book_intro_)
    #         subdoc = create_docx_subdoc(
    #             book_intro_adj,
    #             tn_book.lang_code,
    #             tn_book and tn_book.lang_direction == LangDirEnum.RTL,
    #         )
    #         composer.append(subdoc)
    for bc_book in bc_books:
        # Add the commentary book intro
        subdoc = create_docx_subdoc(bc_book.book_intro, bc_book.lang_code)
        composer.append(subdoc)
    book_codes = {usfm_book.book_code for usfm_book in usfm_books}
    for book_code in book_codes:
        num_chapters = book_chapters[book_code]
        for chapter_num in range(1, num_chapters + 1):
            add_one_column_section(doc)
            # Add chapter intro for each language
            for tn_book in [
                tn_book for tn_book in tn_books if tn_book.book_code == book_code
            ]:
                if chapter_num in tn_book.chapters:
                    subdoc = create_docx_subdoc(
                        chapter_intro(tn_book, chapter_num),
                        tn_book.lang_code,
                        tn_book and tn_book.lang_direction == LangDirEnum.RTL,
                    )
                    composer.append(subdoc)
            for bc_book in [
                bc_book for bc_book in bc_books if bc_book.book_code == book_code
            ]:
                if chapter_num in bc_book.chapters:
                    # Add the chapter commentary.
                    subdoc = create_docx_subdoc(
                        chapter_commentary(bc_book, chapter_num),
                        bc_book.lang_code,
                    )
                    composer.append(subdoc)
            # Add the interleaved USFM chapters
            for usfm_book in [
                usfm_book
                for usfm_book in usfm_books
                if usfm_book.book_code == book_code
            ]:
                if chapter_num in usfm_book.chapters:
                    add_one_column_section(doc)
                    # fmt: off
                    is_rtl = usfm_book and usfm_book.lang_direction == LangDirEnum.RTL
                    # fmt: on
                    subdoc = create_docx_subdoc(
                        usfm_book.chapters[chapter_num].content,
                        usfm_book.lang_code,
                        is_rtl,
                    )
                    composer.append(subdoc)
            # Add the interleaved tn notes
            tn_verses = None
            for tn_book in [
                tn_book for tn_book in tn_books if tn_book.book_code == book_code
            ]:
                tn_verses = tn_chapter_verses(tn_book, chapter_num)
                if tn_verses:
                    add_two_column_section(doc)
                    subdoc = create_docx_subdoc(
                        tn_verses,
                        tn_book.lang_code,
                        tn_book and tn_book.lang_direction == LangDirEnum.RTL,
                    )
                    composer.append(subdoc)
            # Add the interleaved tq questions
            for tq_book in [
                tq_book for tq_book in tq_books if tq_book.book_code == book_code
            ]:
                tq_verses = tq_chapter_verses(tq_book, chapter_num)
                # Add TQ verse content, if any
                if tq_verses:
                    add_two_column_section(doc)
                    subdoc = create_docx_subdoc(
                        tq_verses,
                        tq_book.lang_code,
                        tq_book and tq_book.lang_direction == LangDirEnum.RTL,
                    )
                    composer.append(subdoc)
            add_page_break(doc)
    return composer


def assemble_tn_by_chapter(
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
    book_chapters: Mapping[str, int] = BOOK_CHAPTERS,
) -> Composer:
    """
    Construct the HTML for a 'by chapter' strategy wherein at least
    tn_book_content_units exists.
    """

    def tn_sort_key(resource: TNBook) -> str:
        return resource.lang_code

    def tq_sort_key(resource: TQBook) -> str:
        return resource.lang_code

    def bc_sort_key(resource: BCBook) -> str:
        return resource.lang_code

    tn_books = sorted(tn_books, key=tn_sort_key)
    tq_books = sorted(tq_books, key=tq_sort_key)
    bc_books = sorted(bc_books, key=bc_sort_key)
    doc = Document()
    composer = Composer(doc)
    add_one_column_section(doc)
    # Content team doesn't want TN book intros: https://github.com/WycliffeAssociates/DOC/issues/121
    # for tn_book in tn_books:
    #     if tn_book.book_intro:
    #         book_intro_ = tn_book.book_intro
    #         book_intro_adj = adjust_book_intro_headings(book_intro_)
    #         subdoc = create_docx_subdoc(
    #             book_intro_adj,
    #             tn_book.lang_code,
    #             tn_book and tn_book.lang_direction == LangDirEnum.RTL,
    #         )
    #         composer.append(subdoc)
    for bc_book in bc_books:
        subdoc = create_docx_subdoc(
            bc_book_intro(bc_book),
            bc_book.lang_code,
        )
        composer.append(subdoc)
    book_codes = {tn_book.book_code for tn_book in tn_books}
    for book_code in book_codes:
        num_chapters = book_chapters[book_code]
        for chapter_num in range(1, num_chapters + 1):
            add_one_column_section(doc)
            one_column_html = []
            if chapter_num <= num_chapters:
                one_column_html.append("Chapter {}".format(chapter_num))
            for tn_book in [
                tn_book for tn_book in tn_books if tn_book.book_code == book_code
            ]:
                if chapter_num in tn_book.chapters:
                    # Add the translation notes chapter intro.
                    one_column_html.append(chapter_intro(tn_book, chapter_num))
                    one_column_html_ = "".join(one_column_html)
                    if one_column_html_:
                        subdoc = create_docx_subdoc(
                            one_column_html_,
                            tn_book.lang_code,
                            tn_book and tn_book.lang_direction == LangDirEnum.RTL,
                        )
                        composer.append(subdoc)
            for bc_book in [
                bc_book for bc_book in bc_books if bc_book.book_code == book_code
            ]:
                if chapter_num in bc_book.chapters:
                    # Add the chapter commentary.
                    subdoc = create_docx_subdoc(
                        chapter_commentary(bc_book, chapter_num),
                        bc_book.lang_code,
                    )
                    composer.append(subdoc)
            # Add the interleaved tn notes
            for tn_book in [
                tn_book for tn_book in tn_books if tn_book.book_code == book_code
            ]:
                if chapter_num in tn_book.chapters:
                    tn_verses = tn_chapter_verses(tn_book, chapter_num)
                    if tn_verses:
                        add_two_column_section(doc)
                        subdoc = create_docx_subdoc(
                            tn_verses,
                            tn_book.lang_code,
                            tn_book and tn_book.lang_direction == LangDirEnum.RTL,
                        )
                        composer.append(subdoc)
            # Add the interleaved tq questions
            for tq_book in [
                tq_book for tq_book in tq_books if tq_book.book_code == book_code
            ]:
                tq_verses = tq_chapter_verses(tq_book, chapter_num)
                # Add TQ verse content, if any
                if tq_verses:
                    add_two_column_section(doc)
                    subdoc = create_docx_subdoc(
                        tq_verses,
                        tq_book.lang_code,
                        tq_book and tq_book.lang_direction == LangDirEnum.RTL,
                    )
                    composer.append(subdoc)
            add_page_break(doc)
    return composer


def assemble_tq_by_chapter(
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
    book_chapters: Mapping[str, int] = BOOK_CHAPTERS,
) -> Composer:
    """
    Construct the HTML for a 'by chapter' strategy wherein at least
    tq_book_content_units exists.
    """

    def tq_sort_key(resource: TQBook) -> str:
        return resource.lang_code

    def bc_sort_key(resource: BCBook) -> str:
        return resource.lang_code

    tq_books = sorted(tq_books, key=tq_sort_key)
    bc_books = sorted(bc_books, key=bc_sort_key)
    doc = Document()
    composer = Composer(doc)
    book_codes = {tq_book.book_code for tq_book in tq_books}
    for book_code in book_codes:
        num_chapters = book_chapters[book_code]
        for chapter_num in range(1, num_chapters):
            one_column_html = []
            one_column_html.append("Chapter {}".format(chapter_num))
            for bc_book in [
                bc_book for bc_book in bc_books if bc_book.book_code == book_code
            ]:
                one_column_html.append(chapter_commentary(bc_book, chapter_num))
            if one_column_html:
                add_one_column_section(doc)
                subdoc = create_docx_subdoc("".join(one_column_html), bc_book.lang_code)
                composer.append(subdoc)
            # Add the interleaved tq questions
            for tq_book in [
                tq_book for tq_book in tq_books if tq_book.book_code == book_code
            ]:
                tq_verses = tq_chapter_verses(tq_book, chapter_num)
                if tq_verses:
                    add_two_column_section(doc)
                    subdoc = create_docx_subdoc(
                        tq_verses,
                        tq_book.lang_code,
                        tq_book and tq_book.lang_direction == LangDirEnum.RTL,
                    )
                    composer.append(subdoc)
            add_page_break(doc)
    return composer


# This function could be a little confusing for newcomers. TW lives at
# the language level not the book level, but this function gets invoked
# at the book level due to how the algorithm works. See
# assemble_content_by_book_then_lang above for the conditional that
# invokes it to see the details. At the book level it is almost a noop
# for TW since that is handled elsewhere in
# document_generator.assemble_docx_content.
def assemble_tw_by_chapter(
    usfm_books: Sequence[USFMBook],
    tn_books: Sequence[TNBook],
    tq_books: Sequence[TQBook],
    tw_books: Sequence[TWBook],
    bc_books: Sequence[BCBook],
) -> Composer:
    """Construct the HTML for BC and TW."""
    doc = Document()
    composer = Composer(doc)

    def bc_sort_key(resource: BCBook) -> str:
        return resource.lang_code

    bc_books = sorted(bc_books, key=bc_sort_key)
    for bc_book in bc_books:
        subdoc = create_docx_subdoc(bc_book.book_intro, bc_book.lang_code)
        composer.append(subdoc)
        for chapter in bc_book.chapters.values():
            subdoc = create_docx_subdoc(chapter.commentary, bc_book.lang_code)
            composer.append(subdoc)
            add_page_break(doc)
    return composer
